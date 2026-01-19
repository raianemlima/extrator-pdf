import streamlit as st
import fitz  # PyMuPDF
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
import io
import random
import re
from typing import List, Dict, Tuple
from collections import Counter

# Constantes - Identidade Visual Cursos Duo
COR_VERDE_DUO_RGB = (166, 201, 138)
COR_VERDE_DUO_HEX = "#A6C98A"
COR_VERDE_ESCURO = "#7A9B6E"
COR_VERDE_CLARO = "#D4E7C5"
COR_TEXTO_ESCURO = "#2C3E50"
COR_FUNDO_CLARO = "#F8FCF8"

# Configuração da página
st.set_page_config(
    page_title="Resumo Inteligente - Duo",
    page_icon="🎓",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# CSS customizado avançado - Identidade Visual Profissional
st.markdown(f"""
    <style>
    /* Fonte e cores base */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    html, body, [class*="css"] {{
        font-family: 'Inter', sans-serif;
        color: {COR_TEXTO_ESCURO};
    }}
    
    /* Container principal */
    .main {{
        padding: 1.5rem;
        background: linear-gradient(135deg, #f5f7fa 0%, #f8fcf8 100%);
    }}
    
    /* Cards customizados */
    .card-duo {{
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        border-left: 4px solid {COR_VERDE_DUO_HEX};
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        margin: 1rem 0;
        transition: transform 0.2s;
    }}
    
    .card-duo:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.12);
    }}
    
    /* Badges e tags */
    .badge-duo {{
        background: {COR_VERDE_CLARO};
        color: {COR_VERDE_ESCURO};
        padding: 0.3rem 0.8rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        display: inline-block;
        margin: 0.2rem;
    }}
    
    /* Botões customizados */
    .stDownloadButton > button {{
        background: linear-gradient(135deg, {COR_VERDE_DUO_HEX} 0%, {COR_VERDE_ESCURO} 100%);
        color: white;
        border: none;
        padding: 0.6rem 1.5rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.3s;
    }}
    
    .stDownloadButton > button:hover {{
        transform: scale(1.05);
        box-shadow: 0 4px 12px rgba(166, 201, 138, 0.4);
    }}
    
    /* Responsividade */
    @media (max-width: 768px) {{
        .main {{padding: 0.5rem;}}
        .card-duo {{padding: 1rem;}}
    }}
    
    /* Animações */
    @keyframes slideIn {{
        from {{
            opacity: 0;
            transform: translateY(20px);
        }}
        to {{
            opacity: 1;
            transform: translateY(0);
        }}
    }}
    
    .animated-content {{
        animation: slideIn 0.5s ease-out;
    }}
    </style>
""", unsafe_allow_html=True)


def analisar_conteudo_juridico(texto: str) -> Dict[str, any]:
    """Análise inteligente de conteúdo jurídico."""
    analise = {
        "tema_principal": None,
        "artigos_citados": [],
        "jurisprudencia": [],
        "palavras_chave": [],
        "nivel_complexidade": "Média",
        "tipo_conteudo": "Conceitual"
    }
    
    # Identificação de artigos
    artigos = re.findall(r'art\.?\s*(\d+[A-Z]?(?:-[A-Z])?)', texto, re.IGNORECASE)
    artigos += re.findall(r'artigo\s*(\d+)', texto, re.IGNORECASE)
    analise["artigos_citados"] = list(set(artigos))
    
    # Identificação de jurisprudência
    if any(word in texto.upper() for word in ['STF', 'STJ', 'TST', 'TSE']):
        analise["jurisprudencia"].append("Tribunais Superiores")
    if 'SÚMULA' in texto.upper() or 'SUMULA' in texto.upper():
        sumulas = re.findall(r'[SsúÚ]umula\s*(\d+)', texto)
        analise["jurisprudencia"].extend([f"Súmula {s}" for s in sumulas])
    
    # Identificação de temas
    temas_mapa = {
        "CPI": ["cpi", "comissão parlamentar", "inquérito"],
        "Imunidades": ["imunidade", "inviolabilidade", "prerrogativa"],
        "Processo Legislativo": ["processo legislativo", "emenda", "lei complementar"],
        "Poder Executivo": ["presidente", "vice-presidente", "ministro"],
        "Crime de Responsabilidade": ["impeachment", "crime de responsabilidade"],
        "Garantias Parlamentares": ["parlamentar", "deputado", "senador"],
    }
    
    texto_lower = texto.lower()
    for tema, palavras in temas_mapa.items():
        if any(palavra in texto_lower for palavra in palavras):
            analise["tema_principal"] = tema
            break
    
    # Análise de complexidade
    complexidade_alta = sum([
        len(analise["artigos_citados"]) > 3,
        len(analise["jurisprudencia"]) > 0,
        len(texto.split()) > 100,
    ])
    
    if complexidade_alta >= 3:
        analise["nivel_complexidade"] = "Alta"
    elif complexidade_alta >= 1:
        analise["nivel_complexidade"] = "Média"
    else:
        analise["nivel_complexidade"] = "Básica"
    
    # Tipo de conteúdo
    if analise["jurisprudencia"]:
        analise["tipo_conteudo"] = "Jurisprudencial"
    elif analise["artigos_citados"]:
        analise["tipo_conteudo"] = "Normativo"
    
    return analise


def limpar_texto_total(texto: str) -> str:
    """Limpa e normaliza o texto extraído do PDF."""
    if not texto:
        return ""
    
    # Remove números de rodapé
    texto = re.sub(r'([a-zA-ZáéíóúÁÉÍÓÚçÇ]{3,})(\d+)', r'\1', texto)
    texto = re.sub(r'(\.)(\d+)', r'\1', texto)
    
    # Mapeamento de caracteres especiais
    mapa_sinais = {
        '\u2013': '-', '\u2014': '-', '\u201c': '"', '\u201d': '"',
        '\u2018': "'", '\u2019': "'", '\u2022': '•', '\uf0b7': '•',
        '\uf02d': '-', '\uf0d8': '>', '\u2026': '...', '\u00a0': ' ',
        '? ': '- '
    }
    
    for original, substituto in mapa_sinais.items():
        texto = texto.replace(original, substituto)
    
    return " ".join(texto.split())


def gerar_pergunta_contextualizada(texto: str, analise: Dict = None) -> str:
    """Gera pergunta inteligente baseada no conteúdo."""
    if not analise:
        analise = analisar_conteudo_juridico(texto)
    
    perguntas_tematicas = {
        "CPI": "Quais são os requisitos constitucionais para criação de uma CPI?",
        "Imunidades": "Diferencie imunidade material de imunidade formal dos parlamentares.",
        "Processo Legislativo": "Explique as fases do processo legislativo ordinário.",
        "Crime de Responsabilidade": "Explique o procedimento bifásico do impeachment presidencial.",
    }
    
    if analise["tema_principal"] and analise["tema_principal"] in perguntas_tematicas:
        return perguntas_tematicas[analise["tema_principal"]]
    
    if analise["artigos_citados"]:
        artigo = analise["artigos_citados"][0]
        return f"Qual a importância do art. {artigo} e como ele se aplica ao tema estudado?"
    
    return "Explique os aspectos fundamentais apresentados no material."


def extrair_conteudo_inteligente(texto_completo: str) -> List[Dict[str, any]]:
    """
    Extrai conteúdo relevante do PDF completo de forma inteligente.
    Identifica seções, conceitos-chave e estrutura lógica.
    
    Returns:
        Lista de blocos de conteúdo estruturados e relevantes
    """
    if not texto_completo:
        return []
    
    blocos_relevantes = []
    
    # Limpa o texto
    linhas = [linha.strip() for linha in texto_completo.split('\n') if linha.strip()]
    
    # Padrões que indicam início de seções importantes
    padroes_secao = [
        r'^\d+\.\s+[A-ZÀÁÂÃÉÊÍÓÔÕÚÇ]',  # 1. TITULO
        r'^[A-ZÀÁÂÃÉÊÍÓÔÕÚÇ]{2,}:',      # TITULO:
        r'^[a-z]\)',                      # a) item
        r'^\d+\)',                        # 1) item
        r'^Art\.?\s*\d+',                 # Art. 52
        r'^§\s*\d+',                      # § 1º
        r'^-\s+[A-Z]',                    # - Item
    ]
    
    bloco_atual = []
    titulo_atual = None
    pagina_estimada = 1
    
    for i, linha in enumerate(linhas):
        # Detecta quebra de página (heurística)
        if any(x in linha.lower() for x in ['última atualização', 'dúvidas e sugestões', 'página']):
            pagina_estimada += 1
            continue
        
        # Ignora linhas muito curtas (menos de 20 caracteres)
        if len(linha) < 20:
            # Se é um título potencial, guarda
            if any(re.match(p, linha) for p in padroes_secao):
                titulo_atual = linha
            continue
        
        # Detecta início de nova seção relevante
        is_secao = any(re.match(p, linha) for p in padroes_secao)
        
        if is_secao:
            # Salva bloco anterior se houver conteúdo significativo
            if bloco_atual and len(' '.join(bloco_atual)) > 100:
                texto_bloco = ' '.join(bloco_atual)
                analise = analisar_conteudo_juridico(texto_bloco)
                
                blocos_relevantes.append({
                    'titulo': titulo_atual or 'Conceito',
                    'texto': limpar_texto_total(texto_bloco),
                    'analise': analise,
                    'pag': pagina_estimada,
                    'tipo': 'secao'
                })
            
            # Inicia novo bloco
            bloco_atual = [linha]
            titulo_atual = linha if len(linha) < 100 else None
        else:
            bloco_atual.append(linha)
        
        # Limita tamanho do bloco (máximo 800 palavras)
        if len(' '.join(bloco_atual).split()) > 800:
            texto_bloco = ' '.join(bloco_atual)
            analise = analisar_conteudo_juridico(texto_bloco)
            
            blocos_relevantes.append({
                'titulo': titulo_atual or 'Conceito',
                'texto': limpar_texto_total(texto_bloco),
                'analise': analise,
                'pag': pagina_estimada,
                'tipo': 'secao'
            })
            
            bloco_atual = []
            titulo_atual = None
    
    # Adiciona último bloco
    if bloco_atual and len(' '.join(bloco_atual)) > 100:
        texto_bloco = ' '.join(bloco_atual)
        analise = analisar_conteudo_juridico(texto_bloco)
        
        blocos_relevantes.append({
            'titulo': titulo_atual or 'Conceito',
            'texto': limpar_texto_total(texto_bloco),
            'analise': analise,
            'pag': pagina_estimada,
            'tipo': 'secao'
        })
    
    # Filtra blocos muito similares (remove duplicatas)
    blocos_unicos = []
    textos_vistos = set()
    
    for bloco in blocos_relevantes:
        # Usa primeiras 100 caracteres como fingerprint
        fingerprint = bloco['texto'][:100].lower()
        if fingerprint not in textos_vistos:
            textos_vistos.add(fingerprint)
            blocos_unicos.append(bloco)
    
    return blocos_unicos


def extrair_definicoes_e_conceitos(texto_completo: str) -> List[Dict[str, str]]:
    """
    Extrai definições e conceitos-chave do material.
    Ideal para flashcards de memorização.
    """
    definicoes = []
    
    # Padrões que indicam definições
    padroes_definicao = [
        r'([A-ZÀÁÂÃÉÊÍÓÔÕÚÇ][^.!?]{10,80})\s+(?:é|são|consiste|significa|representa|corresponde)\s+([^.!?]{20,200})[.!?]',
        r'(?:Define-se|Entende-se|Considera-se)\s+([^.!?]{10,80})\s+como\s+([^.!?]{20,200})[.!?]',
        r'([A-ZÀÁÂÃÉÊÍÓÔÕÚÇ][^:]{10,80}):\s*([^.!?]{20,200})[.!?]',
    ]
    
    for padrao in padroes_definicao:
        matches = re.finditer(padrao, texto_completo, re.MULTILINE)
        for match in matches:
            conceito = match.group(1).strip()
            definicao = match.group(2).strip()
            
            # Valida qualidade
            if (len(conceito) > 10 and len(definicao) > 20 and 
                not any(x in conceito.lower() for x in ['página', 'atualização', 'dúvidas'])):
                
                definicoes.append({
                    'conceito': limpar_texto_total(conceito),
                    'definicao': limpar_texto_total(definicao),
                    'tipo': 'definicao'
                })
    
    # Remove duplicatas
    definicoes_unicas = []
    conceitos_vistos = set()
    
    for def_item in definicoes:
        conceito_lower = def_item['conceito'].lower()[:50]
        if conceito_lower not in conceitos_vistos:
            conceitos_vistos.add(conceito_lower)
            definicoes_unicas.append(def_item)
    
    return definicoes_unicas[:30]  # Limita a 30 definições mais relevantes
    """
    Extrai destaques do PDF com análise inteligente E o texto completo.
    
    Returns:
        Tuple: (lista de destaques, texto completo do PDF)
    """
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    highlights = []
    texto_completo = ""
    
    for page_num, page in enumerate(doc):
        # Extrair texto completo da página
        texto_completo += page.get_text() + "\n"
        
        # Extrair destaques
        for annot in page.annots():
            if annot.type[0] == 8:
                texto_extraido = page.get_textbox(annot.rect)
                texto_limpo = limpar_texto_total(texto_extraido)
                
                if texto_limpo:
                    analise = analisar_conteudo_juridico(texto_limpo)
                    
                    highlights.append({
                        "pag": page_num + 1,
                        "texto": texto_limpo,
                        "analise": analise,
                        "timestamp": datetime.now()
                    })
    
    return highlights, texto_completo


def gerar_estatisticas(highlights: List[Dict]) -> Dict:
    """Gera estatísticas sobre o material."""
    if not highlights:
        return {}
    
    stats = {
        "total_itens": len(highlights),
        "temas": Counter(),
        "complexidade": Counter(),
        "tipos_conteudo": Counter(),
        "artigos_mais_citados": Counter(),
    }
    
    for h in highlights:
        if "analise" in h:
            analise = h["analise"]
            
            if analise["tema_principal"]:
                stats["temas"][analise["tema_principal"]] += 1
            
            stats["complexidade"][analise["nivel_complexidade"]] += 1
            stats["tipos_conteudo"][analise["tipo_conteudo"]] += 1
            
            for artigo in analise["artigos_citados"]:
                stats["artigos_mais_citados"][artigo] += 1
    
    return stats


def criar_pdf_resumo(highlights: List[Dict], nome_modulo: str) -> bytes:
    """Cria PDF do resumo."""
    pdf = FPDF()
    pdf.add_page()
    
    # Cabeçalho
    pdf.set_fill_color(*COR_VERDE_DUO_RGB)
    pdf.rect(0, 0, 210, 45, 'F')
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 15, "RESUMO INTELIGENTE", ln=True, align='C')
    pdf.set_font("Helvetica", size=11)
    pdf.cell(0, 8, nome_modulo.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
    pdf.ln(25)
    
    # Conteúdo
    for i, h in enumerate(highlights, 1):
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*COR_VERDE_DUO_RGB)
        pdf.cell(0, 8, f"ITEM {i:02d} | PAG. {h['pag']}", ln=True)
        
        pdf.set_font("Helvetica", size=12)
        pdf.set_text_color(0, 0, 0)
        txt_pdf = h['texto'].encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(190, 7, txt_pdf, align='J')
        pdf.ln(4)
    
    return bytes(pdf.output())


def criar_word_resumo(highlights: List[Dict], nome_modulo: str) -> bytes:
    """Cria documento Word."""
    doc = Document()
    
    # Título
    h_titulo = doc.add_heading(level=0)
    run_titulo = h_titulo.add_run("RESUMO INTELIGENTE")
    run_titulo.font.color.rgb = RGBColor(*COR_VERDE_DUO_RGB)
    
    p_modulo = doc.add_paragraph()
    run_modulo = p_modulo.add_run(nome_modulo)
    run_modulo.bold = True
    
    # Conteúdo
    for i, h in enumerate(highlights, 1):
        p = doc.add_paragraph()
        
        rt = p.add_run(f"ITEM {i:02d} | PÁGINA {h['pag']}\n")
        rt.bold = True
        rt.font.color.rgb = RGBColor(*COR_VERDE_DUO_RGB)
        
        rtx = p.add_run(h['texto'])
        rtx.font.name = 'Arial'
        rtx.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def criar_pdf_perguntas(blocos_conteudo: List[Dict], nome_modulo: str) -> bytes:
    """
    Cria PDF com perguntas e respostas baseado em BLOCOS ESTRUTURADOS.
    Garante coerência e qualidade nas questões.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Cabeçalho do documento
    pdf.set_fill_color(*COR_VERDE_DUO_RGB)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(190, 10, "ROTEIRO DE PERGUNTAS E RESPOSTAS", ln=True, align='C', fill=True)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(190, 6, nome_modulo.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C', fill=True)
    pdf.ln(10)
    
    # Limita quantidade de questões
    blocos_selecionados = blocos_conteudo[:40]
    
    for i, bloco in enumerate(blocos_selecionados, 1):
        analise = bloco.get("analise", {})
        
        # Cabeçalho da questão
        pdf.set_fill_color(248, 252, 248)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(60, 90, 60)
        
        titulo = bloco.get('titulo', 'Conceito')[:60]
        header = f"  QUESTAO {i:02d}: {titulo}"
        
        pdf.cell(190, 8, header.encode('latin-1', 'replace').decode('latin-1'), 
                ln=True, fill=True, border='B')
        
        # Badges informativos
        pdf.set_font("Helvetica", size=8)
        pdf.set_text_color(100, 100, 100)
        badges = []
        
        if analise.get("tema_principal"):
            badges.append(f"Tema: {analise['tema_principal']}")
        if analise.get("nivel_complexidade"):
            badges.append(f"Nivel: {analise['nivel_complexidade']}")
        if bloco.get('pag'):
            badges.append(f"Pag. {bloco['pag']}")
        
        if badges:
            pdf.cell(190, 5, " | ".join(badges).encode('latin-1', 'replace').decode('latin-1'), ln=True)
        
        pdf.ln(2)
        
        # Pergunta contextualizada
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(0, 0, 0)
        pergunta = gerar_pergunta_contextualizada(bloco['texto'], analise)
        pdf.multi_cell(190, 6, f"PERGUNTA: {pergunta}".encode('latin-1', 'replace').decode('latin-1'), align='L')
        
        # Artigos citados
        if analise.get("artigos_citados"):
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(100, 100, 100)
            artigos_str = ", ".join([f"Art. {a}" for a in analise["artigos_citados"][:5]])
            pdf.cell(190, 5, f"Base normativa: {artigos_str}".encode('latin-1', 'replace').decode('latin-1'), ln=True)
        
        pdf.ln(1)
        
        # Resposta
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*COR_VERDE_DUO_RGB)
        pdf.cell(190, 6, "RESPOSTA DO MATERIAL:", ln=True)
        
        pdf.set_font("Helvetica", size=9)
        pdf.set_text_color(20, 20, 20)
        
        # Limita tamanho da resposta para caber melhor
        texto_resposta = bloco['texto']
        if len(texto_resposta) > 600:
            texto_resposta = texto_resposta[:600] + "..."
        
        txt_pr = texto_resposta.encode('latin-1', 'replace').decode('latin-1')
        pdf.set_draw_color(*COR_VERDE_DUO_RGB)
        pdf.multi_cell(190, 5, txt_pr, align='J', border='L')
        
        pdf.ln(5)
    
    return bytes(pdf.output())
    """
    Cria PDF com perguntas e respostas.
    Se texto_completo for fornecido, usa TODO o conteúdo do PDF.
    Senão, usa apenas os destaques.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Se há texto completo, divide em seções relevantes
    if texto_completo:
        # Divide o texto em parágrafos significativos (mais de 100 caracteres)
        paragrafos = [p.strip() for p in texto_completo.split('\n') if len(p.strip()) > 100]
        
        # Limita a um número razoável de questões (máximo 50)
        paragrafos = paragrafos[:50]
        
        for i, paragrafo in enumerate(paragrafos, 1):
            texto_limpo = limpar_texto_total(paragrafo)
            analise = analisar_conteudo_juridico(texto_limpo)
            
            # Cabeçalho
            pdf.set_fill_color(248, 252, 248)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(60, 90, 60)
            
            header = f"  QUESTAO {i:02d}"
            if analise.get("tema_principal"):
                header += f" - {analise['tema_principal']}"
            
            pdf.cell(190, 8, header.encode('latin-1', 'replace').decode('latin-1'), 
                    ln=True, fill=True, border='B')
            
            pdf.ln(2)
            
            # Pergunta
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(0, 0, 0)
            pergunta = gerar_pergunta_contextualizada(texto_limpo, analise)
            pdf.multi_cell(190, 6, f"PERGUNTA: {pergunta}".encode('latin-1', 'replace').decode('latin-1'), align='L')
            
            pdf.ln(1)
            
            # Resposta
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*COR_VERDE_DUO_RGB)
            pdf.cell(190, 6, "RESPOSTA DO MATERIAL:", ln=True)
            
            pdf.set_font("Helvetica", size=10)
            pdf.set_text_color(20, 20, 20)
            txt_pr = texto_limpo[:500].encode('latin-1', 'replace').decode('latin-1')  # Limita tamanho
            pdf.set_draw_color(*COR_VERDE_DUO_RGB)
            pdf.multi_cell(190, 5, txt_pr, align='J', border='L')
            pdf.ln(6)
    else:
        # Usa apenas os destaques (comportamento original)
        for i, h in enumerate(highlights, 1):
            analise = h.get("analise", {})
            
            # Cabeçalho
            pdf.set_fill_color(248, 252, 248)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(60, 90, 60)
            
            header = f"  QUESTAO {i:02d} (Pag. {h['pag']})"
            if analise.get("tema_principal"):
                header += f" - {analise['tema_principal']}"
            
            pdf.cell(190, 8, header.encode('latin-1', 'replace').decode('latin-1'), 
                    ln=True, fill=True, border='B')
            
            pdf.ln(2)
            
            # Pergunta
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(0, 0, 0)
            pergunta = gerar_pergunta_contextualizada(h['texto'], analise)
            pdf.multi_cell(190, 6, f"PERGUNTA: {pergunta}".encode('latin-1', 'replace').decode('latin-1'), align='L')
            
            pdf.ln(1)
            
            # Resposta
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*COR_VERDE_DUO_RGB)
            pdf.cell(190, 6, "RESPOSTA DO MATERIAL:", ln=True)
            
            pdf.set_font("Helvetica", size=10)
            pdf.set_text_color(20, 20, 20)
            txt_pr = h['texto'].encode('latin-1', 'replace').decode('latin-1')
            pdf.set_draw_color(*COR_VERDE_DUO_RGB)
            pdf.multi_cell(190, 5, txt_pr, align='J', border='L')
            pdf.ln(6)
    
    return bytes(pdf.output())


def criar_pdf_flashcards(definicoes: List[Dict], blocos_conteudo: List[Dict], nome_modulo: str) -> bytes:
    """
    Cria PDF com flashcards de alta qualidade.
    Combina definições extraídas + conceitos-chave dos blocos.
    """
    pdf = FPDF()
    pdf.add_page()
    
    # Cabeçalho
    pdf.set_fill_color(*COR_VERDE_DUO_RGB)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(190, 10, "FLASHCARDS DE REVISAO", ln=True, align='C', fill=True)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(190, 6, nome_modulo.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C', fill=True)
    pdf.ln(10)
    
    contador = 0
    
    # Parte 1: Flashcards de DEFINIÇÕES (mais efetivos para memorização)
    if definicoes:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*COR_VERDE_DUO_RGB)
        pdf.cell(190, 8, "PARTE 1: DEFINICOES E CONCEITOS-CHAVE", ln=True)
        pdf.ln(3)
        
        for def_item in definicoes[:25]:  # Máximo 25 definições
            contador += 1
            
            # Frente do card (PERGUNTA)
            pdf.set_fill_color(*COR_VERDE_DUO_RGB)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(190, 8, f" CARD {contador:02d} - FRENTE", border=1, ln=True, fill=True)
            
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "B", 11)
            pergunta = f"O que e/sao: {def_item['conceito']}?"
            pdf.multi_cell(190, 7, pergunta.encode('latin-1', 'replace').decode('latin-1'), 
                          border='LR', align='C')
            
            # Verso do card (RESPOSTA)
            pdf.set_fill_color(240, 240, 240)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*COR_VERDE_DUO_RGB)
            pdf.cell(190, 6, " VERSO - RESPOSTA:", border='LR', ln=True, fill=True)
            
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", size=10)
            resposta = def_item['definicao']
            pdf.multi_cell(190, 6, resposta.encode('latin-1', 'replace').decode('latin-1'), 
                          border='LRB', align='J')
            
            pdf.ln(4)
    
    # Parte 2: Flashcards de CONCEITOS IMPORTANTES (dos blocos estruturados)
    if blocos_conteudo and contador < 50:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(*COR_VERDE_DUO_RGB)
        pdf.cell(190, 8, "PARTE 2: CONCEITOS E REGRAS IMPORTANTES", ln=True)
        pdf.ln(3)
        
        # Seleciona blocos mais relevantes (com artigos, alta complexidade)
        blocos_prioritarios = sorted(
            blocos_conteudo,
            key=lambda x: (
                len(x.get('analise', {}).get('artigos_citados', [])) * 2 +
                (1 if x.get('analise', {}).get('nivel_complexidade') == 'Alta' else 0)
            ),
            reverse=True
        )[:25]
        
        for bloco in blocos_prioritarios:
            if contador >= 50:  # Limite total de cards
                break
            
            contador += 1
            analise = bloco.get('analise', {})
            
            # Frente do card
            pdf.set_fill_color(*COR_VERDE_DUO_RGB)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 10)
            
            titulo_card = bloco.get('titulo', 'Conceito')[:50]
            pdf.cell(190, 8, f" CARD {contador:02d} - {titulo_card}", border=1, ln=True, fill=True)
            
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", "B", 10)
            
            # Pergunta baseada no tema
            pergunta = gerar_pergunta_contextualizada(bloco['texto'], analise)
            pdf.multi_cell(190, 6, pergunta.encode('latin-1', 'replace').decode('latin-1'), 
                          border='LR', align='L')
            
            # Verso do card
            pdf.set_fill_color(240, 240, 240)
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(*COR_VERDE_DUO_RGB)
            pdf.cell(190, 6, " VERSO - RESPOSTA:", border='LR', ln=True, fill=True)
            
            pdf.set_fill_color(255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", size=9)
            
            # Resposta resumida
            texto_card = bloco['texto']
            if len(texto_card) > 400:
                texto_card = texto_card[:400] + "..."
            
            pdf.multi_cell(190, 5, texto_card.encode('latin-1', 'replace').decode('latin-1'), 
                          border='LRB', align='J')
            
            # Artigos relacionados
            if analise.get('artigos_citados'):
                pdf.set_font("Helvetica", "I", 8)
                pdf.set_text_color(100, 100, 100)
                artigos = ", ".join([f"Art. {a}" for a in analise['artigos_citados'][:4]])
                pdf.cell(190, 4, f"Artigos: {artigos}".encode('latin-1', 'replace').decode('latin-1'), 
                        border='LRB', ln=True)
            
            pdf.ln(4)
    
    # Rodapé informativo
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(100, 100, 100)
    pdf.multi_cell(190, 4, 
        "Dica: Imprima frente e verso, recorte e use para revisao espacada. "
        "Leia a FRENTE, tente responder mentalmente, depois confira o VERSO.".encode('latin-1', 'replace').decode('latin-1'))
    
    return bytes(pdf.output())
    """
    Cria PDF com flashcards.
    Se texto_completo for fornecido, usa TODO o conteúdo do PDF.
    Senão, usa apenas os destaques.
    """
    pdf = FPDF()
    pdf.add_page()
    
    # Se há texto completo, cria flashcards de todo o conteúdo
    if texto_completo:
        # Divide em parágrafos significativos
        paragrafos = [p.strip() for p in texto_completo.split('\n') if len(p.strip()) > 100]
        paragrafos = paragrafos[:50]  # Limita quantidade
        
        for i, paragrafo in enumerate(paragrafos, 1):
            texto_limpo = limpar_texto_total(paragrafo)
            
            pdf.set_fill_color(*COR_VERDE_DUO_RGB)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(190, 8, f" CARTAO {i:02d}", border=1, ln=True, fill=True)
            
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", size=11)
            txt_f = texto_limpo[:400].encode('latin-1', 'replace').decode('latin-1')  # Limita tamanho
            pdf.multi_cell(190, 8, txt_f, border=1, align='J')
            pdf.ln(5)
    else:
        # Usa apenas destaques (comportamento original)
        for i, h in enumerate(highlights, 1):
            pdf.set_fill_color(*COR_VERDE_DUO_RGB)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(190, 8, f" CARTAO {i:02d} | PAGINA {h['pag']}", border=1, ln=True, fill=True)
            
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Helvetica", size=11)
            txt_f = h['texto'].encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(190, 8, txt_f, border=1, align='J')
            pdf.ln(5)
    
    return bytes(pdf.output())


def renderizar_cabecalho():
    """Renderiza o cabeçalho."""
    st.markdown(f"""
        <div style="background: linear-gradient(135deg, {COR_VERDE_DUO_HEX} 0%, {COR_VERDE_ESCURO} 100%); 
                    padding: 2rem 1.5rem; border-radius: 15px; text-align: center; 
                    margin-bottom: 2rem; box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
            <h1 style="color: white; margin: 0; font-family: 'Inter', sans-serif; 
                       font-size: 2.2rem; font-weight: 700; letter-spacing: -0.5px;">
                RESUMO INTELIGENTE
            </h1>
            <p style="color: rgba(255,255,255,0.95); margin: 0.8rem 0 0 0; 
                      font-weight: 600; font-size: 1.1rem;">
                Cursos Duo
            </p>
        </div>
    """, unsafe_allow_html=True)


def renderizar_rodape():
    """Renderiza o rodapé."""
    st.markdown("""
        <hr>
        <p style='text-align: center; color: gray; font-size: 0.8em;'>
            Dúvidas: sugestoes@cursosduo.com.br
        </p>
    """, unsafe_allow_html=True)


def main():
    """Função principal."""
    renderizar_cabecalho()
    
    # Upload e configuração
    uploaded_file = st.file_uploader("Suba o material do Cursos Duo (PDF)", type="pdf")
    nome_modulo = st.text_input("Identificação do Material", value="Revisão Ponto 6")
    
    if uploaded_file is None:
        st.info("👆 Faça upload de um PDF com destaques (highlights) para começar.")
        return
    
    try:
        # Extração completa e inteligente
        with st.spinner("🧠 Analisando PDF e extraindo conteúdo inteligente..."):
            highlights, texto_completo, blocos_conteudo, definicoes = extrair_destaques(uploaded_file)
        
        if not highlights:
            st.warning("⚠️ Nenhum destaque encontrado. Marque os trechos importantes.")
            return
        
        # Informações de sucesso
        col_info1, col_info2, col_info3 = st.columns(3)
        with col_info1:
            st.metric("✅ Destaques", len(highlights))
        with col_info2:
            st.metric("📚 Blocos de Conteúdo", len(blocos_conteudo))
        with col_info3:
            st.metric("🎯 Definições Extraídas", len(definicoes))
        
        # Abas
        tab1, tab2, tab3 = st.tabs(["📄 Resumo", "🗂️ Flashcards & P&R", "🧠 Simulado"])
        
        with tab1:
            st.subheader("📄 Resumo Estruturado")
            
            # Filtros
            stats = gerar_estatisticas(highlights)
            col_filtro1, col_filtro2, col_filtro3 = st.columns(3)
            
            with col_filtro1:
                temas_disponiveis = ["Todos"] + [t for t, _ in stats['temas'].most_common()]
                tema_filtro = st.selectbox("🎯 Filtrar por Tema", temas_disponiveis)
            
            with col_filtro2:
                niveis = ["Todos", "Alta", "Média", "Básica"]
                nivel_filtro = st.selectbox("📊 Nível", niveis)
            
            with col_filtro3:
                ordem = st.selectbox("🔢 Ordenar", ["Página", "Complexidade", "Tamanho"])
            
            # Aplicar filtros
            highlights_filtrados = highlights.copy()
            
            if tema_filtro != "Todos":
                highlights_filtrados = [h for h in highlights_filtrados 
                                       if h.get("analise", {}).get("tema_principal") == tema_filtro]
            
            if nivel_filtro != "Todos":
                highlights_filtrados = [h for h in highlights_filtrados 
                                       if h.get("analise", {}).get("nivel_complexidade") == nivel_filtro]
            
            # Ordenação
            if ordem == "Complexidade":
                ordem_complexidade = {"Alta": 3, "Média": 2, "Básica": 1}
                highlights_filtrados.sort(
                    key=lambda x: ordem_complexidade.get(
                        x.get("analise", {}).get("nivel_complexidade", "Média"), 2
                    ), reverse=True
                )
            elif ordem == "Tamanho":
                highlights_filtrados.sort(key=lambda x: len(x["texto"]), reverse=True)
            
            st.info(f"📌 Exibindo {len(highlights_filtrados)} de {len(highlights)} itens")
            
            # Prévia
            with st.expander("👁️ Visualizar prévia", expanded=False):
                for i, h in enumerate(highlights_filtrados[:5], 1):
                    st.markdown(f"**Item {i:02d} | Página {h['pag']}**")
                    st.write(h['texto'][:200] + "..." if len(h['texto']) > 200 else h['texto'])
                    st.divider()
                if len(highlights_filtrados) > 5:
                    st.caption(f"...e mais {len(highlights_filtrados) - 5} itens")
            
            # Downloads
            col1, col2 = st.columns(2)
            
            with col1:
                pdf_resumo = criar_pdf_resumo(highlights_filtrados, nome_modulo)
                st.download_button(
                    "📥 Baixar PDF",
                    pdf_resumo,
                    f"Resumo_{nome_modulo.replace(' ', '_')}.pdf",
                    "application/pdf",
                    use_container_width=True
                )
            
            with col2:
                word_resumo = criar_word_resumo(highlights_filtrados, nome_modulo)
                st.download_button(
                    "📥 Baixar Word",
                    word_resumo,
                    f"Resumo_{nome_modulo.replace(' ', '_')}.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        
        with tab2:
            st.subheader("🗂️ Material de Revisão Avançada")
            
            st.info(f"💡 **Conteúdo Inteligente:** {len(blocos_conteudo)} blocos estruturados + {len(definicoes)} definições extraídas automaticamente!")
            
            # Prévia do conteúdo
            with st.expander("👁️ Ver prévia do conteúdo extraído", expanded=False):
                st.markdown("**📚 Blocos de Conteúdo:**")
                for i, bloco in enumerate(blocos_conteudo[:3], 1):
                    st.markdown(f"**{i}. {bloco.get('titulo', 'Conceito')}**")
                    st.caption(bloco['texto'][:150] + "...")
                    st.divider()
                
                if definicoes:
                    st.markdown("**🎯 Definições Encontradas:**")
                    for i, def_item in enumerate(definicoes[:3], 1):
                        st.markdown(f"**{i}. {def_item['conceito']}**")
                        st.caption(def_item['definicao'][:100] + "...")
                        st.divider()
                
                st.caption(f"...e muito mais conteúdo estruturado!")
            
            col_x, col_y = st.columns(2)
            
            with col_x:
                pdf_perguntas = criar_pdf_perguntas(blocos_conteudo, nome_modulo)
                st.download_button(
                    "📝 Roteiro P&R Completo",
                    pdf_perguntas,
                    f"Roteiro_PR_{nome_modulo.replace(' ', '_')}.pdf",
                    "application/pdf",
                    use_container_width=True,
                    help=f"Contém {len(blocos_conteudo[:40])} questões estruturadas do material completo"
                )
            
            with col_y:
                pdf_flashcards = criar_pdf_flashcards(definicoes, blocos_conteudo, nome_modulo)
                st.download_button(
                    "✂️ Flashcards Inteligentes",
                    pdf_flashcards,
                    f"Flashcards_{nome_modulo.replace(' ', '_')}.pdf",
                    "application/pdf",
                    use_container_width=True,
                    help=f"{len(definicoes)} definições + conceitos-chave em formato frente/verso"
                )
        
        with tab3:
            st.subheader("🧠 Simulado Inteligente - Certo ou Errado")
            
            st.info(f"💡 **Simulado baseado em {len(blocos_conteudo)} blocos de conteúdo estruturado!**")
            
            # Configurações do simulado
            col_config1, col_config2 = st.columns(2)
            
            with col_config1:
                num_questoes = st.slider("Número de questões", 3, min(20, len(blocos_conteudo)), 10)
            
            with col_config2:
                filtro_complexidade = st.selectbox(
                    "Filtrar por complexidade",
                    ["Todas", "Alta", "Média", "Básica"],
                    key="filtro_simulado_complexidade"
                )
            
            # Filtra blocos por complexidade
            blocos_simulado = blocos_conteudo.copy()
            
            if filtro_complexidade != "Todas":
                blocos_simulado = [
                    b for b in blocos_simulado 
                    if b.get('analise', {}).get('nivel_complexidade') == filtro_complexidade
                ]
            
            if len(blocos_simulado) < num_questoes:
                st.warning(f"⚠️ Apenas {len(blocos_simulado)} blocos disponíveis com esta complexidade.")
                num_questoes = len(blocos_simulado)
            
            if not blocos_simulado:
                st.error("❌ Nenhum conteúdo disponível para o simulado.")
                return
            
            # Gera simulado
            if 'simulado_atual' not in st.session_state or st.button("🔄 Gerar Novo Simulado", type="primary"):
                # Seleciona blocos aleatórios
                st.session_state.simulado_atual = random.sample(blocos_simulado, num_questoes)
                st.session_state.respostas_simulado = {}
                st.session_state.gabarito_revelado = False
                st.rerun()
            
            amostra = st.session_state.simulado_atual
            
            # Contador de respostas
            total_respondidas = sum(1 for r in st.session_state.respostas_simulado.values() if r != "Selecione")
            
            st.progress(total_respondidas / len(amostra), 
                       text=f"Progresso: {total_respondidas}/{len(amostra)} questões")
            
            # Questões
            for idx, bloco in enumerate(amostra):
                analise = bloco.get('analise', {})
                
                st.markdown(f"""
                    <div class="card-duo">
                        <strong style="color: {COR_VERDE_ESCURO};">
                            Questão {idx+1} de {len(amostra)}
                        </strong>
                """, unsafe_allow_html=True)
                
                # Badges
                badges = []
                if analise.get("tema_principal"):
                    badges.append(f"🎯 {analise['tema_principal']}")
                if analise.get("nivel_complexidade"):
                    emoji_nivel = {"Alta": "🔥", "Média": "📊", "Básica": "✅"}
                    badges.append(f"{emoji_nivel.get(analise['nivel_complexidade'], '📌')} {analise['nivel_complexidade']}")
                
                if badges:
                    st.markdown(" • ".join(badges))
                
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Exibe texto da questão (limitado para não ficar muito longo)
                texto_questao = bloco['texto']
                if len(texto_questao) > 400:
                    texto_questao = texto_questao[:400] + "..."
                
                st.info(texto_questao)
                
                # Resposta do usuário
                resp = st.radio(
                    "Sua avaliação:",
                    ["Selecione", "Certo", "Errado"],
                    key=f"qz_{idx}",
                    horizontal=True
                )
                
                # Armazena resposta
                if resp != "Selecione":
                    st.session_state.respostas_simulado[idx] = resp
                
                # Mostra feedback se gabarito revelado
                if st.session_state.get('gabarito_revelado', False) and resp != "Selecione":
                    if resp == "Certo":
                        st.success("✅ **Correto!** Esta afirmação está de acordo com o material.")
                        
                        # Dica de revisão
                        if analise.get('artigos_citados'):
                            st.caption(f"💡 Artigos relacionados: {', '.join(analise['artigos_citados'][:3])}")
                    else:
                        st.error("❌ **Errado.** Segundo o material, esta afirmação está correta.")
                        
                        # Mostrar explicação
                        with st.expander("📖 Ver explicação"):
                            st.write(bloco['texto'][:300])
                            if analise.get('artigos_citados'):
                                st.caption(f"📜 Base legal: Artigos {', '.join(analise['artigos_citados'][:3])}")
                
                st.divider()
            
            # Botão para revelar gabarito
            if total_respondidas > 0 and not st.session_state.get('gabarito_revelado', False):
                if st.button("📊 Revelar Gabarito Completo", type="primary", use_container_width=True):
                    st.session_state.gabarito_revelado = True
                    st.rerun()
            
            # Estatísticas finais
            if st.session_state.get('gabarito_revelado', False) and total_respondidas > 0:
                st.markdown("### 🎯 Resultado Final do Simulado")
                
                # Conta acertos (todas as respostas "Certo" são corretas)
                acertos = sum(1 for r in st.session_state.respostas_simulado.values() if r == "Certo")
                percentual = (acertos / total_respondidas) * 100
                
                col_res1, col_res2, col_res3 = st.columns(3)
                
                with col_res1:
                    st.metric("Acertos", f"{acertos}/{total_respondidas}")
                
                with col_res2:
                    st.metric("Aproveitamento", f"{percentual:.1f}%")
                
                with col_res3:
                    if percentual >= 80:
                        st.metric("Status", "🏆 Excelente!")
                    elif percentual >= 60:
                        st.metric("Status", "📈 Bom!")
                    else:
                        st.metric("Status", "📚 Revisar")
                
                # Recomendações
                if percentual < 70:
                    st.warning("💡 **Recomendação:** Revise os flashcards e refaça o simulado para fixar o conteúdo!")
                else:
                    st.success("🎉 **Parabéns!** Você está dominando o conteúdo. Continue praticando!")
        
        renderizar_rodape()
    
    except Exception as e:
        st.error(f"❌ Erro no processamento: {str(e)}")
        st.exception(e)


if __name__ == "__main__":
    main()
